import json
import os
from typing import List, Dict, Any

from PyQt6.QtCore import Qt
from PyQt6.QtWidgets import QMessageBox, QFileDialog
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from libs_auditor.auditor_enum import AuditorKeys
from libs_checker.checker_enum import CheckerKeys
from libs_com.utils_dict import spread_dict
from libs_com.utils_time import get_current_time
from libs_rules.rules_enum import RuleKeys, SeverityLevel
from libs_verifier.verifier_enum import VerifyKeys


def export_json(data: List[Dict[Any, Any]], output_path: str) -> tuple[bool, Exception]:
    """导出为JSON格式"""
    try:
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(spread_dict(data), f, ensure_ascii=False, indent=2)
        return True, None
    except Exception as e:
        return False, e

def export_markdown(data: List[Dict[Any, Any]], output_path: str) -> tuple[bool, Exception]:
    """导出为Markdown格式"""
    try:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("# 漏洞扫描报告\n\n")

            # 按严重程度分组
            severity_groups = group_by_severity_level(data)
            # 生成统计信息
            f.write("## 漏洞统计\n\n")
            for severity, items in severity_groups.items():
                f.write(f"- {severity}: {len(items)}个\n")
            f.write("\n")

            # 生成详细信息
            f.write("## 漏洞详情\n\n")
            for severity, items in severity_groups.items():
                if not items:
                    continue
                f.write(f"### {severity}级别漏洞\n\n")
                for item in items:
                    vuln_type, vuln_locate, vuln_code, vuln_explain, vuln_repair, vuln_status = parse_vuln(item)

                    f.write(f"#### {vuln_type}\n\n")
                    f.write(f"- 漏洞类型: {vuln_type}\n")
                    f.write(f"- 代码位置: {vuln_locate}\n")
                    f.write(f"- 漏洞说明: {vuln_explain}\n")
                    f.write(f"- 漏洞代码:\n```\n{vuln_code}\n```\n")
                    f.write(f"- 修复建议: {vuln_repair}\n")
                    f.write(f"- 审计状态: {vuln_status}\n\n")

        return True, None
    except Exception as e:
        return False, e

def export_html(data: List[Dict[Any, Any]], output_path: str) -> tuple[bool, Exception]:
    """导出为HTML格式"""
    try:
        # HTML模板实现
        html_content = """
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>漏洞扫描报告</title>
            <style>
                body { font-family: Arial, sans-serif; margin: 20px; }
                .vuln-item { margin: 10px 0; padding: 10px; border: 1px solid #ddd; }
                .high { border-left: 5px solid #ff4444; }
                .medium { border-left: 5px solid #ffbb33; }
                .low { border-left: 5px solid #00C851; }
                pre { background-color: #f5f5f5; padding: 10px; }
            </style>
        </head>
        <body>
            <h1>漏洞扫描报告</h1>
        """

        # 按严重程度分组
        severity_groups = group_by_severity_level(data)

        # 添加统计信息
        html_content += "<h2>漏洞统计</h2><ul>"
        for severity, items in severity_groups.items():
            html_content += f"<li>{severity}: {len(items)}个</li>"
        html_content += "</ul>"

        # 添加详细信息
        html_content += "<h2>漏洞详情</h2>"
        for severity, items in severity_groups.items():
            if not items:
                continue
            html_content += f"<h3>{severity}级别漏洞</h3>"
            for item in items:
                vuln_type, vuln_locate, vuln_code, vuln_explain, vuln_repair, vuln_status = parse_vuln(item)

                html_content += f"""
                <div class="vuln-item {severity.lower()}">
                    <h4>{vuln_type}</h4>
                    <p><strong>漏洞类型:</strong> {vuln_type}</p>
                    <p><strong>代码位置:</strong> {vuln_locate}</p>
                    <p><strong>漏洞说明:</strong> {vuln_explain}</p>
                    <p><strong>漏洞代码:</strong></p><pre>{vuln_code}</pre>
                    <p><strong>修复建议:</strong> {vuln_repair}</p>
                    <p><strong>审计状态:</strong> {vuln_status}</p>
                </div>
                """

        html_content += "</body></html>"

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)

        return True, None
    except Exception as e:
        return False, e

def export_docx(data: List[Dict[Any, Any]], output_path: str) -> tuple[bool, Exception]:
    """导出为Word文档格式"""
    try:
        doc = Document()

        # 添加标题
        title = doc.add_heading('漏洞扫描报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加每种漏洞数量的统计信息
        doc.add_heading('漏洞统计信息', level=1)
        severity_groups = group_by_severity_level(data)

        for severity, items in severity_groups.items():
            doc.add_paragraph(f'{severity}: {len(items)}个')

        # 添加详细信息
        doc.add_heading('漏洞详情', level=1)
        for severity, items in severity_groups.items():
            if not items:
                continue

            doc.add_heading(f'漏洞级别: {severity} ', level=2)

            for item in items:
                vuln_type, vuln_locate, vuln_code, vuln_explain, vuln_repair, vuln_status = parse_vuln(item)

                doc_add_issue(doc, idx=None, level=3,
                              vuln_type=vuln_type,
                              vuln_locate=vuln_locate,
                              vuln_level=severity,
                              vuln_affect=vuln_explain,
                              vuln_detail=vuln_code,
                              vuln_repair=vuln_repair
                              )
        doc.save(output_path)
        return True, None
    except Exception as e:
        return False, e



def doc_add_issue(doc, idx, level, vuln_type, vuln_locate, vuln_level, vuln_affect, vuln_detail, vuln_repair):
    # 添加章节标题（level级）
    heading = doc.add_heading(f"{idx}.{vuln_type}", level=level) if idx else doc.add_heading(vuln_type, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 漏洞位置
    p = doc.add_paragraph()
    p.add_run("漏洞位置：").bold = True
    p.add_run(f"{vuln_locate}\n")

    # 漏洞等级
    p.add_run("漏洞等级：").bold = True
    p.add_run(f"{vuln_level}\n")

    # 漏洞影响（即 Abstract）
    p.add_run("漏洞影响：").bold = True
    p.add_run(f"{vuln_affect}\n")

    # 漏洞详情（代码片段）
    p.add_run("漏洞详情：\n").bold = True
    p.add_run(f"{vuln_detail}\n")

    # 修复建议
    p.add_run("修复建议：").bold = True
    p.add_run(f"{vuln_repair}\n")


def parse_vuln(item: Dict):
    original = item[VerifyKeys.ORIGINAL.value]
    parsed = item.get(VerifyKeys.PARSED.value, {})

    vuln_type = original.get(RuleKeys.VULN_TYPE.value)
    vuln_locate = f"{original[CheckerKeys.FILE.value]} ({original[CheckerKeys.LINE.value]}行)"
    vuln_code = original[CheckerKeys.CONTEXT.value]
    vuln_explain = parsed.get(VerifyKeys.EXPLAIN.value) or original.get(RuleKeys.DESCRIPTION.value)
    vuln_repair = parsed.get(VerifyKeys.REPAIR.value) or "待补充"
    vuln_status = item.get(AuditorKeys.AUDITED.value, False)

    return vuln_type, vuln_locate, vuln_code, vuln_explain, vuln_repair, vuln_status


def group_by_severity_level(data: List[Dict[Any, Any]]):
    # 按漏洞等级分组
    severity_groups = {x: [] for x in SeverityLevel.choices()}
    for item in data:
        severity = item.get(VerifyKeys.ORIGINAL.value).get(RuleKeys.SEVERITY.value)
        severity_groups[severity].append(item)

    # 对数据进行排序
    for severity, items in severity_groups.items():
        severity_groups[severity] = sort_vulns_by_vuln_type(items)

    return severity_groups


def sort_vulns_by_vuln_type(items: List[Dict]) -> List[Dict]:
    """对漏洞条目按漏洞类型（VULN_TYPE）进行排序"""
    try:
        return sorted(items, key=lambda x: x[VerifyKeys.ORIGINAL.value][RuleKeys.VULN_TYPE.value])
    except Exception as error:
        # 如果某些条目缺少 VULN_TYPE 字段，可以选择忽略或记录警告
        print(f"sort items by vuln type error: {error}")
        return items  # 或 raise 异常取决于你的需求


def _export_data(parent_ui, export_data, export_type, export_name_prefix='report', base_file=None):
    """通用的导出数据方法"""
    # current_file = main_window.current_file

    try:
        if not export_data:
            QMessageBox.warning(parent_ui, '警告', '没有可导出的数据')
            return

        # 获取保存文件路径
        file_filters = {
            'json': 'JSON Files (*.json)',
            'md': 'Markdown Files (*.md)',
            'html': 'HTML Files (*.html)',
            'docx': 'Word Documents (*.docx)'
        }

        # 检查 base_file 是否为None
        base_file = base_file if base_file else "untitled"
        base_name = os.path.splitext(base_file)[0]
        export_name = f"{base_name}_{get_current_time()}_{export_name_prefix}.{export_type}"
        file_name, _ = QFileDialog.getSaveFileName(parent_ui, '保存报告', export_name, file_filters[export_type])

        if not file_name:
            return

        # 调用相应的导出方法
        export_methods = {
            'json': export_json,
            'md': export_markdown,
            'html': export_html,
            'docx': export_docx
        }

        success, error = export_methods[export_type](export_data, file_name)

        if success:
            QMessageBox.information(parent_ui, '成功', '导出成功')
        else:
            raise error

    except Exception as e:
        QMessageBox.critical(parent_ui, '错误', f'导出失败: {str(e)}')


def export_selected(parent_ui, items, export_type):
    """导出选中的漏洞项"""
    # 收集所有选中的漏洞数据
    selected_data = []

    def collect_item_data(item):
        data = item.data(0, Qt.ItemDataRole.UserRole)
        if data:  # 如果是叶子节点
            selected_data.append(data)
        # 递归处理所有子节点
        for i in range(item.childCount()):
            collect_item_data(item.child(i))

    # 处理所有选中的项
    for item in items:
        collect_item_data(item)
    _export_data(parent_ui, selected_data, export_type, 'selected', parent_ui.current_file)


def export_report(parent_ui, export_type, filtered_data, current_file):
    """导出报告"""
    # 获取当前筛选后的数据
    if filtered_data:
        _export_data(parent_ui, filtered_data, export_type, 'report', current_file)